import io
from typing import Any
from dataclasses import dataclass, field

from abc import ABC, abstractmethod
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from app.analytics.chart_generator import (
    generate_dex_chart,
    generate_gex_chart,
    generate_oi_chart,
    generate_volume_chart,
)


@dataclass
class ReportConfig:
    """Configuration for Word report generation."""
    include_executive_summary: bool = True
    include_asset_data: bool = True
    include_market_interpretation: bool = True
    include_gamma_exposure: bool = True
    include_delta_exposure: bool = True
    include_open_interest: bool = True
    include_volume: bool = True
    include_structural_levels: bool = True
    include_detailed_analysis: bool = True
    include_scenarios: bool = True
    include_conclusions: bool = True
    include_charts: bool = True
    chart_types: list[str] = field(default_factory=lambda: ["gex", "dex", "oi", "volume"])
    language: str = "es"


class SectionGenerator(ABC):
    """Base class for report section generators."""
    
    def __init__(self, doc: Document, config: ReportConfig):
        self.doc = doc
        self.config = config
    
    @abstractmethod
    def add(self, report: dict, analytics_exposure: dict, chart_exposure: dict, 
            chart_source: str | None = None, chart_ratio: float | None = None,
            section_num: int = 1) -> int:
        """Add section to document. Returns next section number."""
        pass
    
    def _format_paragraph(self, p):
        """Apply standard formatting to paragraph."""
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)


class ExecutiveSummaryGenerator(SectionGenerator):
    """Generator for executive summary section."""
    
    def add(self, report: dict, analytics_exposure: dict, chart_exposure: dict,
            chart_source: str | None = None, chart_ratio: float | None = None,
            section_num: int = 1) -> int:
        _add_heading(self.doc, f"{section_num}. Resumen Ejecutivo")
        display_ticker = chart_source if chart_source and chart_source != _g(report, 'ticker', '') else _g(report, 'ticker', 'UNKNOWN')
        p = self.doc.add_paragraph(_executive_summary(report, analytics_exposure, display_ticker))
        self._format_paragraph(p)
        return section_num + 1


class AssetDataGenerator(SectionGenerator):
    """Generator for asset data section."""
    
    def add(self, report: dict, analytics_exposure: dict, chart_exposure: dict,
            chart_source: str | None = None, chart_ratio: float | None = None,
            section_num: int = 1) -> int:
        _add_heading(self.doc, f"{section_num}. Datos del Activo Analizado")
        
        display_ticker = chart_source if chart_source and chart_source != _g(report, 'ticker', '') else _g(report, 'ticker', 'UNKNOWN')
        original_ticker = _g(report, 'ticker', 'UNKNOWN')
        expiration_val = _g(report, 'expiration', 'N/A')
        
        conf_obj = _g(report, 'confidence', None)
        if isinstance(conf_obj, str) or conf_obj is None:
            conf_level = conf_obj or 'N/A'
        else:
            conf_level = _g(conf_obj, 'level', 'N/A')

        _add_kv_table(self.doc, [
            ("Ticker", display_ticker),
            ("Precio Spot", f"${analytics_exposure['spot_price']:.2f}" if analytics_exposure.get('spot_price') is not None else "N/A"),
            ("Vencimiento", expiration_val),
            ("Net GEX", _fmt_money(analytics_exposure.get("net_gamma_exposure"))),
            ("Net DEX", _fmt_money(analytics_exposure.get("net_delta_exposure"))),
            ("Put/Call OI Ratio", f"{analytics_exposure.get('put_call_oi_ratio', 0):.3f}"),
            ("Put/Call Volume Ratio", f"{analytics_exposure.get('put_call_volume_ratio', 0):.3f}"),
            ("Confianza del Análisis", conf_level),
        ])
        
        if chart_source and chart_source != original_ticker:
            ratio_note = f" con ratio {chart_ratio:.2f}" if chart_ratio is not None else ""
            p = self.doc.add_paragraph(
                f"Nota: este informe presenta los niveles en {chart_source}{ratio_note} como equivalencia de {original_ticker}, mientras que la analítica cuantitativa subyacente se realizó sobre {original_ticker}."
            )
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = RGBColor(80, 80, 80)
        
        return section_num + 1


class MarketInterpretationGenerator(SectionGenerator):
    """Generator for market interpretation section."""
    
    def add(self, report: dict, analytics_exposure: dict, chart_exposure: dict,
            chart_source: str | None = None, chart_ratio: float | None = None,
            section_num: int = 1) -> int:
        _add_heading(self.doc, f"{section_num}. Interpretación del Mercado")
        
        ga_desc = _g(_g(report, 'gamma_analysis', {}), 'description', '')
        da_desc = _g(_g(report, 'delta_analysis', {}), 'description', '')
        dealer_desc = _g(_g(report, 'dealer_analysis', {}), 'description', '')
        options_desc = _g(_g(report, 'options_analysis', {}), 'sentiment_description', '')
        
        for text in (ga_desc, da_desc, dealer_desc, options_desc):
            if not text:
                continue
            p = self.doc.add_paragraph(text)
            self._format_paragraph(p)
        
        return section_num + 1


class GammaExposureGenerator(SectionGenerator):
    """Generator for gamma exposure section."""
    
    def add(self, report: dict, analytics_exposure: dict, chart_exposure: dict,
            chart_source: str | None = None, chart_ratio: float | None = None,
            section_num: int = 1) -> int:
        _add_heading(self.doc, f"{section_num}. Gamma Exposure")
        
        original_ticker = _g(report, 'ticker', 'UNKNOWN')
        
        if chart_source and chart_source != original_ticker:
            if chart_ratio is not None:
                note = (
                    f"Los gráficos de este informe están representados en {chart_source} como equivalencia de {original_ticker}, "
                    f"utilizando un ratio de conversión de {chart_ratio:.2f}."
                )
            else:
                note = (
                    f"Los gráficos de este informe están representados en {chart_source} como equivalencia de {original_ticker}."
                )
            p = self.doc.add_paragraph(note)
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = RGBColor(80, 80, 80)
        
        p = self.doc.add_paragraph(
            f"Exposición gamma neta: {_fmt_money(analytics_exposure.get('net_gamma_exposure'))}. "
            f"{_g(_g(report, 'gamma_analysis', {}), 'expected_behavior', '')}"
        )
        self._format_paragraph(p)
        
        if self.config.include_charts and "gex" in self.config.chart_types:
            _add_chart(self.doc, generate_gex_chart(chart_exposure) if generate_gex_chart else None, "Figura 1 — Perfil de Gamma Exposure por Strike")
        
        return section_num + 1


class DeltaExposureGenerator(SectionGenerator):
    """Generator for delta exposure section."""
    
    def add(self, report: dict, analytics_exposure: dict, chart_exposure: dict,
            chart_source: str | None = None, chart_ratio: float | None = None,
            section_num: int = 1) -> int:
        _add_heading(self.doc, f"{section_num}. Delta Exposure")
        
        p = self.doc.add_paragraph(
            f"Exposición delta neta: {_fmt_money(analytics_exposure.get('net_delta_exposure'))}. "
            f"{_g(_g(report, 'delta_analysis', {}), 'hedging_pressure', '')}"
        )
        self._format_paragraph(p)
        
        if self.config.include_charts and "dex" in self.config.chart_types:
            _add_chart(self.doc, generate_dex_chart(chart_exposure) if generate_dex_chart else None, "Figura 2 — Delta Exposure: Calls, Puts y Net Delta por Strike")
        
        return section_num + 1


class OpenInterestGenerator(SectionGenerator):
    """Generator for open interest section."""
    
    def add(self, report: dict, analytics_exposure: dict, chart_exposure: dict,
            chart_source: str | None = None, chart_ratio: float | None = None,
            section_num: int = 1) -> int:
        _add_heading(self.doc, f"{section_num}. Open Interest")
        
        options_liquidity = _g(_g(report, 'options_analysis', {}), 'liquidity_zones', '')
        if options_liquidity:
            p = self.doc.add_paragraph(options_liquidity)
        else:
            p = self.doc.add_paragraph("")
        
        self._format_paragraph(p)
        
        if self.config.include_charts and "oi" in self.config.chart_types:
            _add_chart(self.doc, generate_oi_chart(chart_exposure) if generate_oi_chart else None, "Figura 3 — Distribución de Open Interest por Strike")
        
        return section_num + 1


class VolumeGenerator(SectionGenerator):
    """Generator for volume section."""
    
    def add(self, report: dict, analytics_exposure: dict, chart_exposure: dict,
            chart_source: str | None = None, chart_ratio: float | None = None,
            section_num: int = 1) -> int:
        _add_heading(self.doc, f"{section_num}. Volumen")
        
        p = self.doc.add_paragraph(
            f"Ratio Put/Call de volumen: {analytics_exposure.get('put_call_volume_ratio', 0):.3f}. "
            "Strikes con mayor actividad indican flujo de dinero intradía."
        )
        self._format_paragraph(p)
        
        if self.config.include_charts and "volume" in self.config.chart_types:
            _add_chart(self.doc, generate_volume_chart(chart_exposure) if generate_volume_chart else None, "Figura 4 — Volumen de Calls, Puts y Net Volume por Strike")
        
        return section_num + 1


class StructuralLevelsGenerator(SectionGenerator):
    """Generator for structural levels section."""
    
    def add(self, report: dict, analytics_exposure: dict, chart_exposure: dict,
            chart_source: str | None = None, chart_ratio: float | None = None,
            section_num: int = 1) -> int:
        _add_heading(self.doc, f"{section_num}. Niveles Estructurales")
        
        _add_kv_table(self.doc, [
            ("Put Wall", _fmt_num(analytics_exposure.get("put_wall"))),
            ("Call Wall", _fmt_num(analytics_exposure.get("call_wall"))),
            ("Gamma Flip (Zero Gamma)", _fmt_num(analytics_exposure.get("zero_gamma"))),
            ("Max Pain", _fmt_num(analytics_exposure.get("max_pain"))),
            ("Distancia Spot vs Max Pain",
             _fmt_num(
                 ((analytics_exposure["spot_price"] - analytics_exposure["max_pain"]) / analytics_exposure["spot_price"] * 100)
                 if analytics_exposure.get("max_pain") else None,
                 "%",
             )),
        ])
        
        return section_num + 1


class DetailedAnalysisGenerator(SectionGenerator):
    """Generator for detailed analysis section."""
    
    def add(self, report: dict, analytics_exposure: dict, chart_exposure: dict,
            chart_source: str | None = None, chart_ratio: float | None = None,
            section_num: int = 1) -> int:
        _add_heading(self.doc, f"{section_num}. Análisis Cuantitativo Detallado")
        
        narrative_text = _g(report, 'narrative', '') or ''
        for line in str(narrative_text).split("\n"):
            if line.startswith("# "):
                _add_heading(self.doc, line[2:], level=2)
            elif line.startswith("## "):
                _add_heading(self.doc, line[3:], level=3)
            elif line.startswith("### "):
                _add_heading(self.doc, line[4:], level=4)
            elif line.startswith("- ") or line.startswith("* "):
                p = self.doc.add_paragraph(line[2:], style="List Bullet")
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
            elif line.strip():
                p = self.doc.add_paragraph(line)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
        
        return section_num + 1


class ScenariosGenerator(SectionGenerator):
    """Generator for scenarios section."""
    
    def add(self, report: dict, analytics_exposure: dict, chart_exposure: dict,
            chart_source: str | None = None, chart_ratio: float | None = None,
            section_num: int = 1) -> int:
        _add_heading(self.doc, f"{section_num}. Escenarios Probables")
        
        scenarios_obj = _g(report, "scenarios", {})
        for key in ("principal", "alternative", "risk"):
            sc = _g(scenarios_obj, key, {})
            sc_name = _g(sc, "name", key.capitalize())
            _add_heading(self.doc, sc_name, level=2)
            
            p = self.doc.add_paragraph()
            prob = _g(sc, "probability_pct", None)
            conf = _g(sc, "confidence", "")
            prob_str = f"{prob}%" if prob is not None else "N/A"
            run = p.add_run(f"Probabilidad: {prob_str} | Confianza: {conf}")
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

            narrative_sc = _g(sc, "narrative", "")
            if narrative_sc:
                p = self.doc.add_paragraph(str(narrative_sc))
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)

            p = self.doc.add_paragraph("Factores de soporte:", style="List Bullet")
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            for sf in (_g(sc, "supporting_factors", []) or []):
                p = self.doc.add_paragraph(str(sf), style="List Bullet")
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)

            p = self.doc.add_paragraph("Condiciones de invalidación:", style="List Bullet")
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            for inv in (_g(sc, "invalidation_conditions", []) or []):
                p = self.doc.add_paragraph(str(inv), style="List Bullet")
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
        
        return section_num + 1


class ConclusionsGenerator(SectionGenerator):
    """Generator for conclusions section."""
    
    def add(self, report: dict, analytics_exposure: dict, chart_exposure: dict,
            chart_source: str | None = None, chart_ratio: float | None = None,
            section_num: int = 1) -> int:
        _add_heading(self.doc, f"{section_num}. Conclusiones")
        
        scenarios_obj = _g(report, "scenarios", {})

        p = self.doc.add_paragraph(
            f"Escenario principal ({_g(_g(scenarios_obj, 'principal', {}), 'probability_pct', 'N/A')}): "
            f"{_g(_g(scenarios_obj, 'principal', {}), 'narrative', '')}"
        )
        for run in p.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

        p = self.doc.add_paragraph(
            f"Riesgo de cola ({_g(_g(scenarios_obj, 'risk', {}), 'probability_pct', 'N/A')}): "
            f"{_g(_g(scenarios_obj, 'risk', {}), 'narrative', '')}"
        )
        for run in p.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

        p = self.doc.add_paragraph(
            f"Score de riesgo: {_g(_g(report, 'scores', {}), 'risk_score', 0):.0f}%. "
            f"Volatilidad esperada: {_g(_g(report, 'volatility_analysis', {}), 'description', '')}"
        )
        for run in p.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        return section_num + 1


def _g(obj, key, default=None):
    """Acceso seguro a campos que pueden ser dicts o atributos de objetos."""
    if obj is None:
        return default
    try:
        return obj[key]
    except Exception:
        try:
            return getattr(obj, key)
        except Exception:
            return default


def _fmt_money(value: float | None) -> str:
    if value is None:
        return "N/A"
    abs_v = abs(value)
    if abs_v >= 1_000_000_000:
        return f"${value / 1_000_000_000:.2f}B"
    if abs_v >= 1_000_000:
        return f"${value / 1_000_000:.2f}M"
    if abs_v >= 1_000:
        return f"${value / 1_000:.1f}K"
    return f"${value:.0f}"


def _fmt_num(value: float | None, suffix: str = "") -> str:
    if value is None:
        return "N/A"
    return f"{value:.2f}{suffix}"


def _add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)  # Negro puro
        run.font.bold = True
        run.font.size = Pt(18) if level == 1 else Pt(16) if level == 2 else Pt(14)
    return h


def _add_kv_table(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    
    # Ajustar ancho de columnas
    for row in table.rows:
        row.cells[0].width = Inches(2.5)
        row.cells[1].width = Inches(3.0)
    
    for i, (key, val) in enumerate(rows):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = val
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Negro puro
                    run.font.bold = True

def _add_chart(doc: Document, png_bytes: bytes | None, caption: str):
    doc.add_paragraph()
    if not png_bytes:
        doc.add_paragraph(f"[{caption} — gráfico no disponible en este entorno]")
        return
    stream = io.BytesIO(png_bytes)
    doc.add_picture(stream, width=Inches(6.5))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # Negro puro


def _executive_summary(report: dict, exposure: dict, display_ticker: str | None = None) -> str:
    # Usar acceso seguro que funciona con dicts y objetos
    scores = _g(report, "scores", {})
    gamma = _g(report, "gamma_analysis", {})
    delta = _g(report, "delta_analysis", {})
    dealer = _g(report, "dealer_analysis", {})

    bullish = _g(scores, "bullish_score", 0) or 0
    bearish = _g(scores, "bearish_score", 0) or 0
    bias = "alcista" if bullish > bearish + 10 else "bajista" if bearish > bullish + 10 else "neutral"

    ticker = display_ticker or _g(report, "ticker", "N/A")
    spot = exposure.get("spot_price") if isinstance(exposure, dict) else None
    spot_str = f"{spot:.2f}" if (spot is not None) else "N/A"
    expiration = _g(report, "expiration", "N/A")

    gamma_regime = _g(gamma, "regime_type", "")
    gamma_desc = _g(gamma, "description", "")
    delta_regime = _g(delta, "regime_type", "")
    dealer_regime = _g(dealer, "dealer_gamma_regime", "")
    hedging_impact = _g(dealer, "hedging_impact", "")

    return (
        f"El activo {ticker} cotiza a ${spot_str} con vencimiento {expiration}. El sesgo direccional del análisis es {bias} "
        f"(Bullish {bullish:.0f}% vs Bearish {bearish:.0f}%). "
        f"El régimen gamma es {gamma_regime}: {gamma_desc} "
        f"La presión de cobertura delta es {delta_regime}. "
        f"Los dealers operan en régimen {dealer_regime} gamma con impacto {hedging_impact}. "
        f"Niveles clave: Call Wall {_fmt_num(exposure.get('call_wall'))}, "
        f"Put Wall {_fmt_num(exposure.get('put_wall'))}, "
        f"Gamma Flip {_fmt_num(exposure.get('zero_gamma'))}, "
        f"Max Pain {_fmt_num(exposure.get('max_pain'))}."
    )


def generate_docx_report(
    report: dict,
    analytics_exposure: dict[str, Any],
    chart_exposure: dict[str, Any],
    chart_source: str | None = None,
    chart_ratio: float | None = None,
    config: ReportConfig = None,
) -> io.BytesIO:
    """Generate Word report using modular section generators."""
    if config is None:
        config = ReportConfig()
    
    doc = Document()
    
    # Configurar página
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    
    # Título y metadata
    original_ticker = _g(report, 'ticker', 'UNKNOWN')
    display_ticker = chart_source if chart_source and chart_source != original_ticker else original_ticker
    title = doc.add_heading(f"Reporte de Inteligencia de Mercado — {display_ticker}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    expiration_val = _g(report, 'expiration', 'N/A')
    fetched_val = _g(report, 'fetched_at', '')
    fetched_str = fetched_val[:19] if isinstance(fetched_val, str) else str(fetched_val)
    run = subtitle.add_run(f"Vencimiento: {expiration_val}  |  Generado: {fetched_str} UTC")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Build generators based on config
    generators = []
    if config.include_executive_summary:
        generators.append(ExecutiveSummaryGenerator(doc, config))
    if config.include_asset_data:
        generators.append(AssetDataGenerator(doc, config))
    if config.include_market_interpretation:
        generators.append(MarketInterpretationGenerator(doc, config))
    if config.include_gamma_exposure:
        generators.append(GammaExposureGenerator(doc, config))
    if config.include_delta_exposure:
        generators.append(DeltaExposureGenerator(doc, config))
    if config.include_open_interest:
        generators.append(OpenInterestGenerator(doc, config))
    if config.include_volume:
        generators.append(VolumeGenerator(doc, config))
    if config.include_structural_levels:
        generators.append(StructuralLevelsGenerator(doc, config))
    if config.include_detailed_analysis:
        generators.append(DetailedAnalysisGenerator(doc, config))
    if config.include_scenarios:
        generators.append(ScenariosGenerator(doc, config))
    if config.include_conclusions:
        generators.append(ConclusionsGenerator(doc, config))

    # Execute generators sequentially
    section_num = 1
    for generator in generators:
        section_num = generator.add(
            report=report,
            analytics_exposure=analytics_exposure,
            chart_exposure=chart_exposure,
            chart_source=chart_source,
            chart_ratio=chart_ratio,
            section_num=section_num
        )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
